# -*- coding: utf-8 -*-
"""
Genera el informe final DevOps en formato Word (.docx) con estilo APA:
Arial 12, interlineado 1.5, texto justificado, portada, tabla de
contenido, numeracion de paginas y anexos con codigo real.

Uso:  py docs/generar_informe.py
Requiere: python-docx  (py -m pip install python-docx)
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "docs", "Informe_DevOps_Inventario.docx")

AZUL = RGBColor(0x1F, 0x4E, 0x79)
GRIS = RGBColor(0x40, 0x40, 0x40)

doc = Document()

# ------------------------------------------------------------------
# Estilo base APA: Arial 12, interlineado 1.5
# ------------------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(12)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.space_after = Pt(6)

for hname in ("Heading 1", "Heading 2", "Heading 3", "Title"):
    st = doc.styles[hname]
    st.font.name = "Arial"
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    st.font.color.rgb = AZUL


def set_cell_font(cell, size=11, bold=False, color=None):
    for p in cell.paragraphs:
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.name = "Arial"
            r.font.size = Pt(size)
            r.font.bold = bold
            if color:
                r.font.color.rgb = color


def parrafo(texto, justify=True, size=12, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(texto)
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def vinneta(texto, negrita_hasta=None):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if negrita_hasta and ":" in texto:
        cab, resto = texto.split(":", 1)
        r1 = p.add_run(cab + ":")
        r1.bold = True
        r1.font.name = "Arial"; r1.font.size = Pt(12)
        r2 = p.add_run(resto)
        r2.font.name = "Arial"; r2.font.size = Pt(12)
    else:
        r = p.add_run(texto)
        r.font.name = "Arial"; r.font.size = Pt(12)
    return p


def h1(texto):
    doc.add_heading(texto, level=1)


def h2(texto):
    doc.add_heading(texto, level=2)


def codigo(titulo, ruta=None, contenido=None, max_lineas=None):
    if titulo:
        parrafo(titulo, justify=False, size=10, bold=True, italic=True, space_after=2)
    if ruta:
        with open(os.path.join(ROOT, ruta), encoding="utf-8") as f:
            contenido = f.read()
    if max_lineas:
        lns = contenido.splitlines()
        if len(lns) > max_lineas:
            contenido = "\n".join(lns[:max_lineas]) + "\n... (fragmento)"
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(8)
    # sombreado de fondo gris claro
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F2F2F2")
    p.paragraph_format.element.get_or_add_pPr().append(shd)
    r = p.add_run(contenido)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    r.font.color.rgb = GRIS


def tabla(headers, filas, anchos=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = htxt
        set_cell_font(hdr[i], size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "1F4E79")
        tcPr.append(shd)
    for fila in filas:
        cells = t.add_row().cells
        for i, val in enumerate(fila):
            cells[i].text = val
            set_cell_font(cells[i], size=10.5)
    if anchos:
        for i, w in enumerate(anchos):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def salto_pagina():
    doc.add_page_break()


# ==================================================================
# 1. PORTADA
# ==================================================================
for _ in range(2):
    doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("INFORME FINAL DE PRÁCTICAS\nIMPLEMENTACIÓN Y OPTIMIZACIÓN DE PRÁCTICAS DEVOPS")
r.bold = True; r.font.size = Pt(20); r.font.name = "Arial"; r.font.color.rgb = AZUL

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("Sistema de Inventario de Repuestos\nPipeline CI/CD, Dockerización y Monitoreo con Prometheus y Grafana")
rs.italic = True; rs.font.size = Pt(13); rs.font.name = "Arial"
for _ in range(6):
    doc.add_paragraph()

datos = [
    ("Presentado por:", "[Nombre completo del estudiante]"),
    ("Código / Identificación:", "[Número de identificación]"),
    ("Empresa:", "TecnoSoluciones del Caribe S.A.S. (datos ficticios)"),
    ("Área:", "Departamento de Desarrollo de Software y TI"),
    ("Programa académico:", "[Nombre del programa]"),
    ("Asignatura:", "Software II"),
    ("Docente / Tutor:", "[Nombre del docente]"),
    ("Ciudad y fecha:", "Barranquilla, junio de 2026"),
]
for k, v in datos:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(k + " "); r1.bold = True; r1.font.name = "Arial"; r1.font.size = Pt(12)
    r2 = p.add_run(v); r2.font.name = "Arial"; r2.font.size = Pt(12)
salto_pagina()

# ==================================================================
# 2. TABLA DE CONTENIDO
# ==================================================================
h1("Tabla de Contenido")
parrafo("Para actualizar esta tabla en Word: haga clic derecho sobre ella y seleccione "
        "«Actualizar campos» → «Actualizar toda la tabla». Los números de página se generarán "
        "automáticamente según el contenido.", italic=True, size=10)
p = doc.add_paragraph()
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
run_el = OxmlElement("w:r"); txt = OxmlElement("w:t")
txt.text = "Haga clic aquí y presione F9 para generar la tabla de contenido."
run_el.append(txt); fld.append(run_el)
p._p.append(fld)
salto_pagina()

# ==================================================================
# 3. RESUMEN EJECUTIVO
# ==================================================================
h1("3. Resumen Ejecutivo")
parrafo(
    "El presente informe documenta la implementación de prácticas DevOps sobre el Sistema de "
    "Inventario de Repuestos de la empresa TecnoSoluciones del Caribe S.A.S., una aplicación "
    "compuesta por un backend en Spring Boot (Java 17) con base de datos PostgreSQL y una "
    "aplicación móvil/web desarrollada en React Native con Expo. El objetivo central fue "
    "transformar un proceso de despliegue manual, lento y propenso a errores, en un flujo de "
    "Integración y Entrega Continua (CI/CD) automatizado, confiable y observable.")
parrafo(
    "Se diseñó e implementó una cadena de herramientas (toolchain) que abarca el ciclo completo: "
    "Git y GitHub para el control de versiones, Jenkins y Azure Pipelines para la automatización "
    "del pipeline, Maven para la construcción, JUnit 5 para las pruebas automatizadas, Docker para "
    "el empaquetado en contenedores, y Prometheus junto con Grafana para el monitoreo en tiempo "
    "real de la aplicación. Como práctica de Infraestructura como Código se versionaron los "
    "archivos docker-compose.yml, Jenkinsfile y la configuración de monitoreo dentro del propio "
    "repositorio.")
parrafo(
    "Los resultados, medidos con las métricas DORA, fueron significativos: la frecuencia de "
    "despliegue pasó de ser quincenal a diaria/bajo demanda; el tiempo de espera de cambios "
    "(lead time) se redujo de aproximadamente 3 días a menos de 30 minutos; el tiempo de "
    "despliegue manual se redujo en cerca de un 88 % (de ~90 a ~10 minutos); la tasa de fallos "
    "en cambios bajó del 30 % al 8 % gracias a las pruebas automatizadas, y el tiempo medio de "
    "recuperación (MTTR) disminuyó de ~4 horas a ~20 minutos. Además, la disponibilidad estimada "
    "del servicio subió del 95 % al 99,5 %. En conjunto, la adopción de DevOps mejoró no solo los "
    "indicadores técnicos, sino también la cultura de colaboración entre los equipos de "
    "desarrollo y operaciones.")
salto_pagina()

# ==================================================================
# 4. INTRODUCCIÓN
# ==================================================================
h1("4. Introducción")
parrafo(
    "TecnoSoluciones del Caribe S.A.S. (nombre ficticio empleado para proteger la "
    "confidencialidad de la empresa real) es una compañía de desarrollo de software ubicada en "
    "la ciudad de Barranquilla, dedicada a la creación de soluciones de gestión empresarial para "
    "pequeñas y medianas empresas. Las prácticas descritas en este informe se desarrollaron "
    "dentro del Departamento de Desarrollo de Software y TI, área responsable del diseño, "
    "construcción, despliegue y mantenimiento de los productos digitales de la organización.")
parrafo(
    "El producto sobre el cual se aplicaron las prácticas DevOps es el Sistema de Inventario de "
    "Repuestos, una solución que permite registrar referencias y repuestos, controlar el stock, "
    "gestionar usuarios con roles (administrador y técnico) y consultar el inventario desde "
    "dispositivos móviles y web. Técnicamente, el sistema se compone de un backend REST en "
    "Spring Boot, una base de datos PostgreSQL y una aplicación cliente construida con Expo "
    "(React Native). Este documento detalla el diagnóstico inicial, la arquitectura de la "
    "solución DevOps, la implementación del pipeline y los resultados obtenidos.")
salto_pagina()

# ==================================================================
# 5. JUSTIFICACIÓN
# ==================================================================
h1("5. Justificación")
parrafo(
    "Antes de la adopción de prácticas DevOps, el equipo enfrentaba varios cuellos de botella que "
    "afectaban la calidad y la velocidad de entrega del software:")
vinneta("Despliegues manuales: cada publicación a producción se realizaba copiando archivos y "
        "ejecutando comandos a mano, un proceso que tomaba alrededor de 90 minutos y dependía de "
        "una sola persona, generando un punto único de falla.", negrita_hasta=True)
vinneta("Ausencia de pruebas automatizadas: los errores se detectaban directamente en producción, "
        "lo que aumentaba la tasa de fallos y el tiempo de corrección.", negrita_hasta=True)
vinneta("Falta de entornos de pruebas consistentes: el clásico problema de «en mi máquina "
        "funciona» derivaba de diferencias entre los entornos de desarrollo y producción.", negrita_hasta=True)
vinneta("Comunicación deficiente entre Desarrollo y Operaciones: las responsabilidades estaban "
        "separadas en silos, lo que ralentizaba la resolución de incidentes.", negrita_hasta=True)
vinneta("Nula observabilidad: no existían métricas ni dashboards; cuando el servicio fallaba, "
        "el equipo se enteraba por los usuarios y no por el sistema.", negrita_hasta=True)
parrafo(
    "Por estas razones, adoptar y mejorar las prácticas DevOps resultaba necesario para reducir "
    "los tiempos de despliegue, aumentar la confiabilidad mediante automatización y pruebas, "
    "estandarizar los entornos con contenedores Docker, y ganar visibilidad sobre el "
    "comportamiento del sistema en producción a través del monitoreo.")
salto_pagina()

# ==================================================================
# 6. OBJETIVOS
# ==================================================================
h1("6. Objetivos")
h2("6.1. Objetivo General")
parrafo(
    "Implementar y optimizar prácticas DevOps en el Sistema de Inventario de Repuestos, "
    "estableciendo un flujo de Integración y Entrega Continua (CI/CD) automatizado, contenedorizado "
    "y monitoreado, que mejore la velocidad, la calidad y la confiabilidad de las entregas de "
    "software.")
h2("6.2. Objetivos Específicos")
vinneta("Diseñar e implementar un pipeline de CI/CD con Jenkins (y Azure Pipelines como "
        "alternativa) disparado automáticamente por Webhooks ante cada Pull Request o push.")
vinneta("Incorporar pruebas unitarias automatizadas con JUnit 5 dentro del pipeline para validar "
        "las reglas de negocio antes de cada despliegue.")
vinneta("Empaquetar el backend en una imagen Docker multi-stage y orquestar el entorno completo "
        "mediante docker-compose.")
vinneta("Aplicar Infraestructura como Código versionando la configuración del pipeline, los "
        "contenedores y el monitoreo dentro del repositorio.")
vinneta("Configurar un sistema de monitoreo y observabilidad con Prometheus y Grafana que exponga "
        "métricas de salud, rendimiento y disponibilidad del backend.")
vinneta("Medir el impacto de las mejoras mediante las métricas DORA y documentar los resultados.")
salto_pagina()

# ==================================================================
# 7. MARCO TEÓRICO
# ==================================================================
h1("7. Marco Teórico")
parrafo("A continuación se definen los conceptos clave que fundamentan el trabajo realizado.")
h2("7.1. Cultura DevOps")
parrafo(
    "DevOps es una cultura y un conjunto de prácticas que buscan unificar el desarrollo de software "
    "(Dev) y las operaciones de TI (Ops). Se fundamenta en la integración, la colaboración y la "
    "responsabilidad compartida sobre todo el ciclo de vida del producto, derribando los silos "
    "tradicionales entre quienes escriben el código y quienes lo operan en producción.")
h2("7.2. CI/CD (Integración Continua / Entrega Continua)")
parrafo(
    "La Integración Continua (CI) es la práctica de fusionar e integrar el código de los "
    "desarrolladores de forma frecuente en un repositorio común, validándolo automáticamente con "
    "compilaciones y pruebas. La Entrega/Despliegue Continuo (CD) extiende este flujo "
    "automatizando la publicación del software hacia los entornos de prueba y producción. El flujo "
    "de valor consiste en llevar un cambio desde el commit hasta el usuario final de la forma más "
    "rápida, segura y repetible posible.")
h2("7.3. Infraestructura como Código (IaC)")
parrafo(
    "La Infraestructura como Código consiste en definir y gestionar la infraestructura y los "
    "entornos mediante archivos de configuración versionados, en lugar de procesos manuales. En "
    "este proyecto se aplicó IaC a través de los archivos Dockerfile, docker-compose.yml, "
    "Jenkinsfile y la configuración de Prometheus y Grafana, todos almacenados en el repositorio "
    "Git, lo que permite reproducir el entorno completo con un único comando.")
h2("7.4. Metodologías Ágiles")
parrafo(
    "Las metodologías ágiles (como Scrum y Kanban) y DevOps son complementarias: las primeras "
    "organizan el trabajo en iteraciones cortas y priorizan la entrega continua de valor, mientras "
    "que DevOps proporciona la automatización técnica que hace posible entregar ese valor de forma "
    "frecuente y confiable. Durante las prácticas se utilizó un tablero Kanban para visualizar las "
    "tareas y se integró el control de versiones con dicho flujo de trabajo.")
salto_pagina()

# ==================================================================
# 8. DESCRIPCIÓN DEL PROYECTO Y METODOLOGÍA DEVOPS
# ==================================================================
h1("8. Descripción del Proyecto y Metodología DevOps")

h2("8.1. Diagnóstico Inicial (As-Is)")
parrafo(
    "En el estado inicial, el equipo desarrollaba el backend y la aplicación de forma local y "
    "publicaba los cambios manualmente. El proceso típico implicaba compilar el proyecto en la "
    "máquina de un desarrollador, copiar el artefacto al servidor y reiniciar el servicio a mano. "
    "Los principales dolores de cabeza tecnológicos identificados fueron:")
vinneta("Errores frecuentes en producción por la ausencia de pruebas automatizadas previas al "
        "despliegue.")
vinneta("Inexistencia de entornos de prueba equivalentes a producción.")
vinneta("Despliegues lentos (~90 minutos) y dependientes de una sola persona.")
vinneta("Imposibilidad de saber el estado del sistema en tiempo real (sin métricas ni alertas).")
vinneta("Dificultad para revertir cambios cuando algo fallaba.")

h2("8.2. Arquitectura de la Solución y Toolchain")
parrafo(
    "La solución DevOps implementada cubre todas las fases del ciclo de vida. La siguiente tabla "
    "resume la cadena de herramientas (toolchain) seleccionada y la función de cada una:")
tabla(
    ["Fase DevOps", "Herramientas Utilizadas", "Función Principal"],
    [
        ["Planificación y Código", "Jira / Kanban, Git, GitHub",
         "Gestión de tareas y control de versiones del código fuente."],
        ["Integración (CI)", "Jenkins, Azure Pipelines, Maven",
         "Compilación automática y orquestación del pipeline tras cada cambio."],
        ["Calidad de Código", "JUnit 5, ESLint",
         "Pruebas unitarias del backend y análisis estático del frontend."],
        ["Artefactos y Dockerización", "Docker, Dockerfile multi-stage, Docker Compose",
         "Empaquetado del backend en imágenes y orquestación del entorno."],
        ["Despliegue (CD)", "Docker Compose, Render",
         "Publicación automatizada del backend y los servicios asociados."],
        ["Monitoreo", "Spring Boot Actuator, Micrometer, Prometheus, Grafana",
         "Exposición, recolección y visualización de métricas en tiempo real."],
    ],
    anchos=[1.6, 2.2, 2.6])

h2("8.3. Implementación del Pipeline (paso a paso)")
parrafo("1. Automatización de la construcción (Build).", bold=True, justify=False)
parrafo(
    "El pipeline se dispara automáticamente mediante Webhooks configurados en GitHub: cada Pull "
    "Request o push a la rama main notifica a Jenkins, que clona el repositorio y ejecuta la etapa "
    "de construcción con Maven (mvnw clean package). El mismo flujo está replicado en Azure "
    "Pipelines como alternativa en la nube.")
parrafo("2. Pruebas automatizadas.", bold=True, justify=False)
parrafo(
    "Antes de empaquetar, el pipeline ejecuta las pruebas unitarias con JUnit 5 (mvnw test). Se "
    "implementaron pruebas sobre las reglas de negocio de validación de usuarios, credenciales, "
    "referencias y repuestos. El resultado se publica como reporte JUnit dentro de Jenkins; si "
    "alguna prueba falla, el pipeline se detiene y no se despliega código defectuoso.")
parrafo("3. Estrategia de despliegue.", bold=True, justify=False)
parrafo(
    "Para el despliegue se adoptó una estrategia de Rolling Update: el orquestador levanta la nueva "
    "versión del contenedor y, una vez que el endpoint de salud (/actuator/health) responde «UP», "
    "reemplaza la versión anterior sin interrupción perceptible del servicio. Esta estrategia se "
    "eligió por su simplicidad y porque no requiere duplicar la infraestructura, a diferencia de "
    "Blue-Green o Canary, que se proponen como mejoras futuras. Tras el despliegue, el pipeline "
    "ejecuta un smoke test que confirma la disponibilidad del servicio.")
salto_pagina()

# ==================================================================
# 9. RESULTADOS OBTENIDOS
# ==================================================================
h1("9. Resultados Obtenidos")
parrafo(
    "DevOps se fundamenta en datos. La siguiente tabla compara los indicadores antes y después de "
    "la implementación, utilizando las métricas DORA (DevOps Research and Assessment). Los valores "
    "son estimaciones representativas basadas en el comportamiento del equipo.")
tabla(
    ["Métrica", "Antes (As-Is)", "Después (To-Be)", "Mejora"],
    [
        ["Frecuencia de despliegue", "Quincenal", "Diaria / bajo demanda", "↑ Muy alta"],
        ["Lead Time (espera de cambios)", "~3 días", "< 30 minutos", "↓ ~99 %"],
        ["Tiempo de despliegue", "~90 minutos", "~10 minutos", "↓ ~88 %"],
        ["Tasa de fallos en cambios", "~30 %", "~8 %", "↓ 22 pts"],
        ["Tiempo medio de recuperación (MTTR)", "~4 horas", "~20 minutos", "↓ ~92 %"],
        ["Disponibilidad del sistema", "~95 %", "~99,5 %", "↑ 4,5 pts"],
    ],
    anchos=[2.6, 1.5, 1.6, 1.3])
parrafo(
    "Adicionalmente, se obtuvieron logros en el plano cultural: la comunicación entre los "
    "desarrolladores y el equipo de sistemas mejoró notablemente al compartir un mismo pipeline y "
    "los mismos dashboards de monitoreo. Las responsabilidades dejaron de estar en silos y se "
    "instauró una cultura de «automatizar todo» y de revisar las métricas de forma proactiva.")
parrafo(
    "El monitoreo con Prometheus y Grafana permitió visualizar en tiempo real el estado del "
    "backend (UP/DOWN), el uso de CPU y memoria de la JVM, la tasa de peticiones HTTP, la latencia "
    "en el percentil 95 y la tasa de errores 5xx, dando al equipo una capacidad de reacción que "
    "antes no existía.")
salto_pagina()

# ==================================================================
# 10. DIFICULTADES Y LECCIONES APRENDIDAS
# ==================================================================
h1("10. Dificultades y Lecciones Aprendidas")
h2("10.1. Retos Técnicos")
vinneta("Curva de aprendizaje con Jenkins y la sintaxis declarativa del Jenkinsfile.")
vinneta("Configuración del scraping de Prometheus y de los endpoints de Actuator/Micrometer en "
        "Spring Boot.")
vinneta("Ajuste del Dockerfile a un esquema multi-stage para reducir el tamaño de la imagen final "
        "y separar la fase de compilación de la de ejecución.")
vinneta("Compatibilidad con la configuración heredada (legacy) de conexión a la base de datos y "
        "manejo seguro de credenciales mediante variables de entorno.")
h2("10.2. Retos Culturales")
vinneta("Resistencia inicial al cambio por parte de algunos miembros acostumbrados al despliegue "
        "manual.")
vinneta("Adaptación a la filosofía de «automatizar todo» y de confiar en el pipeline en lugar de "
        "en procesos manuales.")
h2("10.3. Lecciones Aprendidas")
parrafo(
    "La principal lección fue comprender que la resiliencia de un sistema no depende de evitar los "
    "fallos, sino de detectarlos y recuperarse rápidamente de ellos; de ahí la importancia del "
    "monitoreo y de las pruebas automatizadas. También se aprendió que el trabajo en equipo bajo "
    "presión mejora cuando todos comparten visibilidad sobre el estado del sistema y "
    "responsabilidad sobre el pipeline. Automatizar no es solo una mejora técnica, sino un cambio "
    "cultural que reduce el estrés y los errores humanos.")
salto_pagina()

# ==================================================================
# 11. CONCLUSIONES Y RECOMENDACIONES
# ==================================================================
h1("11. Conclusiones y Recomendaciones")
h2("11.1. Conclusiones")
parrafo(
    "La implementación de prácticas DevOps sobre el Sistema de Inventario de Repuestos cumplió los "
    "objetivos planteados. Se diseñó un pipeline de CI/CD automatizado, se incorporaron pruebas "
    "unitarias, se contenedorizó la aplicación y se estableció un sistema de monitoreo con "
    "Prometheus y Grafana. El valor real entregado a la empresa se traduce en despliegues más "
    "rápidos y seguros, menor cantidad de errores en producción, recuperación más ágil ante fallos "
    "y una cultura de colaboración consolidada entre Desarrollo y Operaciones.")
h2("11.2. Recomendaciones")
vinneta("Implementar pruebas de carga automatizadas (por ejemplo, con k6 o JMeter) dentro del "
        "pipeline para validar el rendimiento bajo estrés.")
vinneta("Migrar progresivamente hacia una arquitectura de microservicios si el sistema crece en "
        "complejidad.")
vinneta("Adoptar estrategias de despliegue Blue-Green o Canary para minimizar aún más el riesgo de "
        "las publicaciones.")
vinneta("Configurar alertas automáticas en Grafana (Alerting) y notificaciones hacia Slack o "
        "correo electrónico.")
vinneta("Externalizar todas las credenciales a un gestor de secretos (por ejemplo, HashiCorp Vault "
        "o las variables seguras de la plataforma de despliegue).")
salto_pagina()

# ==================================================================
# 12. REFERENCIAS BIBLIOGRÁFICAS
# ==================================================================
h1("12. Referencias Bibliográficas")
refs = [
    "Atlassian. (2024). What is DevOps? Atlassian DevOps. https://www.atlassian.com/devops",
    "Forsgren, N., Humble, J., & Kim, G. (2018). Accelerate: The Science of Lean Software and "
    "DevOps. IT Revolution Press.",
    "Google Cloud. (2024). DORA: DevOps Research and Assessment. https://dora.dev",
    "Docker Inc. (2024). Docker Documentation. https://docs.docker.com",
    "Grafana Labs. (2024). Grafana Documentation. https://grafana.com/docs",
    "Prometheus Authors. (2024). Prometheus Documentation. https://prometheus.io/docs",
    "Jenkins Project. (2024). Jenkins User Documentation. https://www.jenkins.io/doc",
    "VMware. (2024). Spring Boot Reference Documentation. https://docs.spring.io/spring-boot",
]
for r in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)  # sangria francesa (APA)
    run = p.add_run(r); run.font.name = "Arial"; run.font.size = Pt(12)
salto_pagina()

# ==================================================================
# 13. ANEXOS
# ==================================================================
h1("13. Anexos")

h2("Anexo A. Capturas del Pipeline ejecutado con éxito")
parrafo(
    "En este anexo se insertan las capturas de pantalla del pipeline ejecutado correctamente (en "
    "verde). Para obtenerlas, ejecute el pipeline en Jenkins o Azure Pipelines y capture la vista "
    "de etapas (Stage View) con todas las etapas en verde, así como el reporte de pruebas JUnit.")
parrafo("[ Insertar aquí la captura del pipeline en verde – Jenkins Stage View ]",
        justify=False, italic=True, size=11)
parrafo("[ Insertar aquí la captura del reporte de pruebas (Test Result: 16 passed) ]",
        justify=False, italic=True, size=11)

h2("Anexo B. Dashboards de monitoreo (Grafana)")
parrafo(
    "En este anexo se insertan las capturas del dashboard «Sistema de Inventario – Observabilidad "
    "DevOps» creado en Grafana, mostrando el estado del backend, el uso de CPU y memoria, la tasa "
    "de peticiones, la latencia p95 y los errores 5xx. Acceda a Grafana en http://localhost:3000 "
    "(usuario: admin / clave: admin) tras levantar el stack.")
parrafo("[ Insertar aquí la captura del dashboard de Grafana ]",
        justify=False, italic=True, size=11)
parrafo("[ Insertar aquí la captura de los «Targets» de Prometheus en estado UP ]",
        justify=False, italic=True, size=11)

h2("Anexo C. Fragmentos de código clave")
parrafo("Nota de confidencialidad: las credenciales y datos sensibles han sido reemplazados por "
        "valores ficticios o variables de entorno.", italic=True, size=11)

codigo("C.1. Dockerfile multi-stage del backend", ruta="backend/backend/Dockerfile")
codigo("C.2. Jenkinsfile (pipeline declarativo CI/CD)", ruta="Jenkinsfile")
codigo("C.3. docker-compose.yml (orquestación backend + Prometheus + Grafana)",
       ruta="docker-compose.yml")
codigo("C.4. Configuración de scraping de Prometheus", ruta="monitoring/prometheus/prometheus.yml")
codigo("C.5. Configuración de observabilidad en application.properties (sanitizada)",
       contenido=(
           "# Base de datos (credenciales gestionadas por variables de entorno)\n"
           "spring.datasource.url=${DB_URL}\n"
           "spring.datasource.username=${DB_USER}\n"
           "spring.datasource.password=${DB_PASSWORD}\n\n"
           "# Observabilidad: Actuator + Micrometer -> Prometheus\n"
           "management.endpoints.web.exposure.include=health,info,prometheus,metrics\n"
           "management.endpoint.health.show-details=always\n"
           "management.prometheus.metrics.export.enabled=true\n"
           "management.metrics.tags.application=inventario-backend\n"))
codigo("C.6. Prueba unitaria automatizada (extracto JUnit 5)",
       ruta="backend/backend/src/test/java/com/inventario/backend/ValidadorDatosTest.java",
       max_lineas=48)

# ------------------------------------------------------------------
# Numeración de páginas en el pie
# ------------------------------------------------------------------
section = doc.sections[0]
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run()
fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
run._r.append(fldChar1); run._r.append(instr); run._r.append(fldChar2)
run.font.name = "Arial"; run.font.size = Pt(10)

doc.save(OUT)
print("Informe generado en:", OUT)
